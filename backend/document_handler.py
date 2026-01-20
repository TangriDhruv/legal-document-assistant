# backend/document_handler.py
"""
Document processing module using python-docx
IMPROVED: Keeps type inference, no circular dependencies
"""

import re
from typing import List, Dict
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX


def infer_placeholder_type(placeholder_name: str, context: str) -> str:
    """
    Infer placeholder type from name and context
    Returns: text, currency, date, person_name, company_name, address, email, phone
    
    No LLM calls - uses keyword matching on context
    """
    
    combined_text = (placeholder_name + " " + context).lower()
    
    # Type-specific keyword patterns
    type_patterns = {
        'currency': [
            'amount', 'price', 'cost', 'fee', 'payment', 'paid', 'invest', 'purchase',
            'dollar', 'salary', 'sum', 'total', 'value', 'rate', '$', 'thousand'
        ],
        'date': [
            'date', 'when', 'day', 'month', 'year', 'time', 'period', 'until', 'from',
            'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august',
            'september', 'october', 'november', 'december', '20', '2024', '2025'
        ],
        'person_name': [
            'investor', 'founder', 'director', 'officer', 'partner', 'representative',
            'mr', 'ms', 'dr', 'attorney', 'signatory', 'member', 'person', "person's"
        ],
        'company_name': [
            'company', 'corporation', 'corp', 'inc', 'llc', 'ltd', 'lp', 'entity',
            'organization', 'business', 'firm', 'group', 'enterprise', 'venture'
        ],
        'address': [
            'address', 'street', 'city', 'state', 'zip', 'zipcode', 'road', 'avenue',
            'boulevard', 'drive', 'lane', 'location', 'place', 'building'
        ],
        'email': [
            'email', 'mail', 'contact', '@', '.com', '.org', '.net', '.edu'
        ],
        'phone': [
            'phone', 'number', 'telephone', 'mobile', 'cell', 'contact', '(', ')'
        ]
    }
    
    # Count keyword matches for each type
    type_scores = {type_name: 0 for type_name in type_patterns.keys()}
    
    for type_name, keywords in type_patterns.items():
        for keyword in keywords:
            if keyword in combined_text:
                type_scores[type_name] += 1
    
    # Get type with highest score
    if max(type_scores.values()) > 0:
        best_type = max(type_scores, key=type_scores.get)
        return best_type
    
    # Default to text if no keywords matched
    return 'text'


def extract_document_text(docx_path: str) -> str:
    """
    Extract plain text from .docx file
    
    Args:
        docx_path: Path to .docx file
        
    Returns:
        Plain text from all paragraphs
    """
    doc = Document(docx_path)
    text_parts = []
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    
    return "\n".join(text_parts)


def find_placeholders(text: str) -> List[Dict]:
    """
    Find all [Bracketed] placeholders in document text
    Infers type from context using keyword analysis
    
    Args:
        text: Document text to search
        
    Returns:
        List of placeholder dictionaries with name, context, type, etc.
    """
    # Pattern to match [text] or [text with spaces]
    pattern = r'\[([^\]]+)\]'
    matches = re.finditer(pattern, text)
    
    placeholders = []
    seen = set()
    
    for match in matches:
        placeholder_name = match.group(1).strip()
        
        # Skip empty brackets
        if not placeholder_name:
            continue
        
        # Skip duplicates (case-insensitive)
        placeholder_name_lower = placeholder_name.lower()
        if placeholder_name_lower in seen:
            print(f"⏭️ Skipping duplicate placeholder: [{placeholder_name}]")
            continue
        seen.add(placeholder_name_lower)
        
        start, end = match.span()
        
        # Extract context (150 chars before and after for type inference)
        context_start = max(0, start - 150)
        context_end = min(len(text), end + 150)
        context = text[context_start:context_end].strip()
        
        # INFER TYPE from context using keyword matching
        inferred_type = infer_placeholder_type(placeholder_name, context)
        
        # Generate description based on inferred type
        type_descriptions = {
            'currency': 'Amount in dollars/currency',
            'date': 'Date (e.g., MM/DD/YYYY or text format)',
            'person_name': 'Person\'s full name',
            'company_name': 'Company or organization name',
            'address': 'Full address',
            'email': 'Email address',
            'phone': 'Phone number',
            'text': 'Text value'
        }
        
        description = type_descriptions.get(inferred_type, f'Please provide: {placeholder_name.lower()}')
        
        placeholder_dict = {
            "name": placeholder_name,  # EXACT name as it appears in document
            "context": context,
            "filled": False,
            "value": None,
            "type": inferred_type,  # INFERRED from context
            "description": description
        }
        
        placeholders.append(placeholder_dict)
        print(f"✓ Found placeholder: [{placeholder_name}] (Type: {inferred_type})")
    
    return placeholders


def fill_placeholders(
    docx_path: str, 
    values: Dict[str, str], 
    output_path: str
) -> None:
    """
    Replace [Placeholder] with actual values in .docx
    Preserves original formatting
    
    Args:
        docx_path: Path to original .docx
        values: Dictionary of placeholder -> value mappings
                Keys MUST match exact placeholder names from document
        output_path: Path for completed .docx
    """
    
    print(f"\n{'='*70}")
    print(f"📄 FILL_PLACEHOLDERS starting...")
    print(f"Input file: {docx_path}")
    print(f"Values to fill: {values}")
    print(f"Output file: {output_path}")
    print(f"{'='*70}\n")
    
    doc = Document(docx_path)
    
    # Track what we've filled
    replacements_made = {placeholder: False for placeholder in values.keys()}
    
    def replace_in_paragraph(paragraph, values_dict):
        """Replace placeholders in a paragraph, handling split runs"""
        
        # Get combined text to check what's in this paragraph
        full_text = paragraph.text
        
        # Check which placeholders are in this paragraph
        placeholders_here = [
            (name, val) for name, val in values_dict.items() 
            if f"[{name}]" in full_text
        ]
        
        if not placeholders_here:
            return
        
        # Combine all run texts
        combined_text = "".join([run.text for run in paragraph.runs])
        
        # Replace all placeholders in the combined text
        for placeholder_name, replacement_value in placeholders_here:
            placeholder_str = f"[{placeholder_name}]"
            if placeholder_str in combined_text:
                combined_text = combined_text.replace(
                    placeholder_str, 
                    replacement_value or ""
                )
                replacements_made[placeholder_name] = True
                print(f"  ✓ Replaced [{placeholder_name}] in paragraph")
        
        # Clear existing runs and set new text
        for run in paragraph.runs:
            run.text = ""
        
        if paragraph.runs:
            paragraph.runs[0].text = combined_text
        else:
            paragraph.add_run(combined_text)
    
    # Replace in paragraphs
    print("🔄 Processing paragraphs...")
    for i, paragraph in enumerate(doc.paragraphs):
        if any(f"[{name}]" in paragraph.text for name in values.keys()):
            print(f"  Found placeholders in paragraph {i}")
            replace_in_paragraph(paragraph, values)
    
    # Replace in tables
    print("🔄 Processing tables...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if any(f"[{name}]" in paragraph.text for name in values.keys()):
                        print(f"  Found placeholders in table[{table_idx}] row[{row_idx}] cell[{cell_idx}] para[{para_idx}]")
                        replace_in_paragraph(paragraph, values)
    
    # Save document
    doc.save(output_path)
    
    # Print summary
    print(f"\n{'='*70}")
    print(f"✅ Document saved to: {output_path}")
    print(f"\nReplacement summary:")
    for placeholder_name, was_replaced in replacements_made.items():
        status = "✓ REPLACED" if was_replaced else "✗ NOT FOUND"
        print(f"  [{placeholder_name}]: {status}")
    print(f"{'='*70}\n")
    
    # Warn if nothing was replaced
    if not any(replacements_made.values()):
        print("⚠️  WARNING: No placeholders were replaced!")
        print("   Check that placeholder names in values match document exactly.")